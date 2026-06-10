"""Documents services (ORM port of the NestJS documents module).

Ports three NestJS providers into this single module:

  - ``DocumentsService``           -> template CRUD + form-page generation + Outlook draft
  - ``DocumentGenerationService``  -> chat-driven (LLM) document generation
  - ``DocumentExportService``      -> markdown -> .pdf / .docx export

Raw ``pg`` queries become SQLAlchemy ORM queries. Cross-module calls (graph, rag)
use LAZY imports inside the methods to avoid the rag/ingestion/query import cycle at
boot. Audit logging is written directly via the ``AuditLog`` ORM model (the audit
table is owned by the ``common`` module, which exposes no service yet). OpenAI is
called through the ``openai`` python SDK (``AsyncOpenAI``).
"""

import json
import re
from typing import Any

from sqlalchemy import delete, select
from sqlalchemy.ext.asyncio import AsyncSession

from app import logger
from app.core.constants import LLM_MODEL
from app.core.settings import settings
from app.core.utils import normalize_keys
from app.models.audit import AuditLog
from app.models.documents import DocumentTemplate as DocumentTemplateModel
from app.prompts import (
    DOCUMENT_EXTRACTION_PROMPT,
    DOCUMENT_GENERATION_PROMPT,
    DOCUMENT_LAYOUT_PROMPT,
    GROUNDED_DOCUMENT_PROMPT,
    build_document_extraction_user_prompt,
    build_document_generation_user_prompt,
    build_document_layout_user_prompt,
    build_grounded_document_user_prompt,
)
from app.schema.auth import AuthContext
from app.schema.common import (
    Citation,
    DocumentDraft,
    DocumentTemplate,
    DocumentTemplateType,
    RagQueryResponse,
    RetrievedChunk,
    TemplateVariable,
    VectorSource,
)

# Columns selected for a DocumentTemplate row (mirrors TEMPLATE_COLUMNS in the
# NestJS service: id, organization_id, workspace_id, name, type, content, variables,
# is_default, created_at, updated_at).
_TEMPLATE_COLUMNS = (
    DocumentTemplateModel.id,
    DocumentTemplateModel.organization_id,
    DocumentTemplateModel.workspace_id,
    DocumentTemplateModel.name,
    DocumentTemplateModel.type,
    DocumentTemplateModel.content,
    DocumentTemplateModel.variables,
    DocumentTemplateModel.is_default,
    DocumentTemplateModel.created_at,
    DocumentTemplateModel.updated_at,
)

# Phrases that imply the user wants the document grounded in their M365 data.
_SOURCE_REFERENCE = re.compile(
    r"\b(sharepoint|onedrive|outlook|email|e-mail|mail|inbox|document|file|invoice|"
    r"report|contract|attachment|pdf|from\s+the|based\s+on|go\s+over)\b",
    re.IGNORECASE,
)

_PLACEHOLDER_RE = re.compile(r"\{\{\s*([\w.]+)\s*\}\}")


class DocumentError(Exception):
    """Base error -> maps to HTTP responses in the router (parity with Nest exceptions)."""


class DocumentNotFoundError(DocumentError):
    """Raised when a template is missing (NestJS NotFoundException -> 404)."""


class DocumentBadRequestError(DocumentError):
    """Raised on invalid input (NestJS BadRequestException -> 400)."""


def extract_placeholders(content: str) -> list[str]:
    """Pull every distinct ``{{placeholder}}`` from a template body, in first-seen order."""
    seen: set[str] = set()
    out: list[str] = []
    for match in _PLACEHOLDER_RE.finditer(content):
        key = match.group(1)
        if key not in seen:
            seen.add(key)
            out.append(key)
    return out


def _escape_regexp(s: str) -> str:
    return re.sub(r"([.*+?^${}()|\[\]\\])", r"\\\1", s)


def _to_template(row: Any) -> DocumentTemplate:
    """Map a selected ORM row tuple to the shared :class:`DocumentTemplate` schema."""
    raw_vars = row.variables or []
    variables = [TemplateVariable.model_validate(v) for v in raw_vars]
    return DocumentTemplate(
        id=str(row.id),
        organization_id=str(row.organization_id),
        workspace_id=str(row.workspace_id) if row.workspace_id else None,
        name=row.name,
        type=DocumentTemplateType(row.type.value if hasattr(row.type, "value") else row.type),
        content=row.content,
        variables=variables,
        is_default=row.is_default,
        created_at=row.created_at.isoformat() if row.created_at else "",
        updated_at=row.updated_at.isoformat() if row.updated_at else "",
    )


class DocumentsService:
    """Template CRUD, form-page generation and Outlook drafting."""

    def __init__(self, db: AsyncSession) -> None:
        self.db = db
        self._openai: Any | None = None

    # --- OpenAI client (lazy) -------------------------------------------------

    def _client(self) -> Any:
        if self._openai is None:
            if not settings.openai_api_key:
                msg = "OPENAI_API_KEY is not configured"
                raise RuntimeError(msg)
            from openai import AsyncOpenAI

            self._openai = AsyncOpenAI(api_key=settings.openai_api_key)
        return self._openai

    async def _log_audit(
        self,
        ctx: AuthContext,
        action: str,
        resource_type: str | None = None,
        resource_id: str | None = None,
        metadata: dict[str, Any] | None = None,
    ) -> None:
        """Insert an audit_logs row (port of common AuditService.logAudit)."""
        self.db.add(
            AuditLog(
                organization_id=ctx.organization_id,
                user_id=ctx.user_id,
                action=action,
                resource_type=resource_type,
                resource_id=resource_id,
                metadata_=metadata or {},
            )
        )
        await self.db.commit()

    # --- Template CRUD --------------------------------------------------------

    async def list_templates(self, ctx: AuthContext) -> list[DocumentTemplate]:
        result = await self.db.execute(
            select(*_TEMPLATE_COLUMNS)
            .where(DocumentTemplateModel.organization_id == ctx.organization_id)
            .order_by(DocumentTemplateModel.type.asc(), DocumentTemplateModel.name.asc())
        )
        return [_to_template(row) for row in result.all()]

    async def get_template(self, ctx: AuthContext, id: str) -> DocumentTemplate:
        result = await self.db.execute(
            select(*_TEMPLATE_COLUMNS).where(
                DocumentTemplateModel.id == id,
                DocumentTemplateModel.organization_id == ctx.organization_id,
            )
        )
        row = result.first()
        if row is None:
            raise DocumentNotFoundError("Template not found")
        return _to_template(row)

    async def create_template(
        self, ctx: AuthContext, input: "CreateTemplateInput"
    ) -> DocumentTemplate:
        name = (input.get("name") or "").strip()
        content = (input.get("content") or "").strip()
        if not name:
            raise DocumentBadRequestError("Template name is required")
        if not content:
            raise DocumentBadRequestError("Template content is required")

        # Derive variables from the {{placeholders}} in the body when not supplied.
        supplied = input.get("variables")
        if supplied:
            variables = [
                v.model_dump(exclude_none=True) if isinstance(v, TemplateVariable) else v
                for v in supplied
            ]
        else:
            variables = [{"key": key, "label": key} for key in extract_placeholders(content)]

        model = DocumentTemplateModel(
            organization_id=ctx.organization_id,
            workspace_id=input.get("workspace_id"),
            user_id=ctx.user_id,
            name=name,
            type=input["type"].value
            if isinstance(input["type"], DocumentTemplateType)
            else input["type"],
            content=content,
            variables=variables,
            is_default=False,
        )
        self.db.add(model)
        await self.db.commit()
        await self.db.refresh(model)

        await self._log_audit(
            ctx, "document.template.create", "template", str(model.id)
        )
        return await self.get_template(ctx, str(model.id))

    async def update_template(
        self, ctx: AuthContext, id: str, input: "UpdateTemplateInput"
    ) -> DocumentTemplate:
        await self.get_template(ctx, id)  # ownership / existence check

        result = await self.db.execute(
            select(DocumentTemplateModel).where(
                DocumentTemplateModel.id == id,
                DocumentTemplateModel.organization_id == ctx.organization_id,
            )
        )
        model = result.scalar_one_or_none()
        if model is None:
            raise DocumentNotFoundError("Template not found")

        # COALESCE($n, col): only overwrite when a (trimmed) value is supplied.
        name = input.get("name")
        if name is not None and name.strip():
            model.name = name.strip()
        type_ = input.get("type")
        if type_ is not None:
            model.type = type_.value if isinstance(type_, DocumentTemplateType) else type_
        content = input.get("content")
        if content is not None and content.strip():
            model.content = content.strip()
        if "variables" in input and input.get("variables") is not None:
            model.variables = [
                v.model_dump(exclude_none=True) if isinstance(v, TemplateVariable) else v
                for v in input["variables"]
            ]

        await self.db.commit()

        await self._log_audit(ctx, "document.template.update", "template", id)
        return await self.get_template(ctx, id)

    async def delete_template(self, ctx: AuthContext, id: str) -> None:
        result = await self.db.execute(
            delete(DocumentTemplateModel).where(
                DocumentTemplateModel.id == id,
                DocumentTemplateModel.organization_id == ctx.organization_id,
            )
        )
        await self.db.commit()
        if result.rowcount == 0:
            raise DocumentNotFoundError("Template not found")
        await self._log_audit(ctx, "document.template.delete", "template", id)

    async def find_template(
        self,
        ctx: AuthContext,
        type: DocumentTemplateType | None = None,
        name_hint: str | None = None,
    ) -> DocumentTemplate | None:
        """Resolve the best template: explicit name match, then type default, then any."""
        templates = await self.list_templates(ctx)
        if not templates:
            return None

        hint = (name_hint or "").strip().lower()
        if hint:
            by_name = next(
                (t for t in templates if hint in t.name.lower()), None
            )
            if by_name:
                return by_name

        if type:
            of_type = [t for t in templates if t.type == type]
            if of_type:
                return next((t for t in of_type if t.is_default), of_type[0])

        return None

    # --- Generation -----------------------------------------------------------

    async def render_template(
        self,
        template: DocumentTemplate,
        variables: dict[str, str],
        context_text: str | None = None,
    ) -> str:
        """Fill a template with provided values and let the LLM polish the result."""
        content = template.content
        for key, value in variables.items():
            if value is None or str(value).strip() == "":
                continue
            content = re.sub(
                rf"\{{\{{\s*{_escape_regexp(key)}\s*\}}\}}",
                str(value),
                content,
            )

        missing_keys = extract_placeholders(content)

        client = self._client()
        completion = await client.chat.completions.create(
            model=LLM_MODEL,
            messages=[
                {"role": "system", "content": DOCUMENT_GENERATION_PROMPT},
                {
                    "role": "user",
                    "content": build_document_generation_user_prompt(
                        content, variables, missing_keys, context_text
                    ),
                },
            ],
        )
        message = completion.choices[0].message.content if completion.choices else None
        return (message or "").strip() or content

    async def generate(
        self, ctx: AuthContext, req: "GenerateDocumentRequestInput"
    ) -> dict[str, Any]:
        """Form-page entry point: generate from a template id + supplied variables."""
        template = await self.get_template(ctx, req["template_id"])
        generated = await self.render_template(template, req.get("variables") or {})
        await self._log_audit(
            ctx,
            "document.generate",
            "template",
            req["template_id"],
            {"type": template.type.value},
        )
        return {
            "content": generated,
            "template_id": req["template_id"],
            "type": template.type,
        }

    # --- Outlook draft --------------------------------------------------------

    async def save_email_draft(
        self, ctx: AuthContext, draft: dict[str, Any]
    ) -> dict[str, str]:
        if not (draft.get("body") or "").strip():
            raise DocumentBadRequestError("Draft body is required")

        # LAZY import: graph is a sibling module (avoids circular imports at boot).
        from app.services.graph_service import GraphService

        graph = GraphService(self.db)
        result = await graph.create_draft(
            ctx,
            {
                "subject": (draft.get("subject") or "").strip() or "(no subject)",
                "body": draft["body"],
                "to": draft.get("to"),
            },
        )
        await self._log_audit(ctx, "document.email_drafted", "mail", result["id"])
        return result


class DocumentGenerationService:
    """Chat-driven (LLM) document generation grounded in M365 / project content."""

    def __init__(self, db: AsyncSession) -> None:
        self.db = db
        self.documents = DocumentsService(db)
        self._openai: Any | None = None

    def _client(self) -> Any:
        if self._openai is None:
            if not settings.openai_api_key:
                msg = "OPENAI_API_KEY is not configured"
                raise RuntimeError(msg)
            from openai import AsyncOpenAI

            self._openai = AsyncOpenAI(api_key=settings.openai_api_key)
        return self._openai

    async def generate_from_prompt(
        self,
        ctx: AuthContext,
        request: dict[str, Any],
        options: dict[str, Any] | None = None,
    ) -> RagQueryResponse:
        try:
            return await self._run_generation(ctx, request, options)
        except Exception as err:
            message = str(err)
            logger.error(f"[docgen] generation failed: {message}")
            hint = (
                " It looks like the document_templates table is not migrated — run "
                "`pnpm db:migrate` and restart the API."
                if re.search(r"is_default|column|relation|does not exist", message, re.IGNORECASE)
                else ""
            )
            return self._plain_response(
                f"I couldn't generate that document right now.{hint}"
            )

    async def _run_generation(
        self,
        ctx: AuthContext,
        request: dict[str, Any],
        options: dict[str, Any] | None = None,
    ) -> RagQueryResponse:
        # Callers pass option dicts with camelCase keys (projectId/workspaceId);
        # normalize to the snake_case keys this code reads.
        options = normalize_keys(options or {})
        templates = await self.documents.list_templates(ctx)

        if not request["description"].strip():
            return self._help_response(templates, "Tell me what to generate.")
        if not templates:
            return self._plain_response(
                "There are no document templates set up yet. An admin can add one on "
                "the Documents page (or seed the defaults), then I can draft estimates, "
                "job summaries, and customer emails for you."
            )

        # Auto-detect: pull supporting M365 content when the request references it
        # (a project chat is always grounded in its own knowledge base).
        context = await self._maybe_retrieve(ctx, request["description"], options)

        extracted = await self._extract(
            request["description"], templates, context["context_text"]
        )

        template = await self.documents.find_template(
            ctx, type=extracted["type"], name_hint=extracted["name_hint"]
        )

        # No matching template, but we have source content -> free-form grounded doc.
        if template is None:
            if context["has_content"]:
                return await self._generate_grounded(ctx, request, context)
            return self._help_response(
                templates, "I couldn't match that to one of your templates."
            )

        variables = dict(extracted["variables"])
        if (
            template.type == DocumentTemplateType.customer_email
            and extracted["subject"]
            and "subject" not in variables
        ):
            variables["subject"] = extracted["subject"]

        generated = await self.documents.render_template(
            template, variables, context["context_text"] or None
        )

        await self.documents._log_audit(
            ctx,
            "document.generate.chat",
            "template",
            template.id,
            {
                "type": template.type.value,
                "explicit": request.get("explicit"),
                "grounded": context["has_content"],
            },
        )

        return self._draft_response(template, generated, extracted, context["citations"])

    async def _maybe_retrieve(
        self,
        ctx: AuthContext,
        description: str,
        options: dict[str, Any] | None = None,
    ) -> dict[str, Any]:
        """Retrieve M365 / project context when the request references source content."""
        options = options or {}
        is_project = bool(options.get("project_id"))
        if not is_project and not _SOURCE_REFERENCE.search(description):
            return {"context_text": "", "citations": [], "has_content": False}

        try:
            # LAZY import: rag is a sibling in the import cycle.
            from app.services.rag_service import RagService

            rag = RagService(self.db)
            result = await rag.query(
                ctx,
                description,
                {
                    "project_id": options.get("project_id"),
                    "workspace_id": options.get("workspace_id"),
                    "sources": None if is_project else _detect_sources(description),
                    "top_k": 6,
                    "bypass_topic_guard": True,
                    "project_only": is_project,
                },
            )
            chunks = result.chunks or []
            context_text = "\n\n---\n\n".join(
                f"[{i + 1}] ({_chunk_label(c)})\n{c.content}"
                for i, c in enumerate(chunks)
            )
            return {
                "context_text": context_text,
                "citations": result.citations or [],
                "has_content": len(chunks) > 0,
            }
        except Exception as err:
            logger.warning(f"[docgen] context retrieval failed: {err}")
            return {"context_text": "", "citations": [], "has_content": False}

    async def _extract(
        self,
        description: str,
        templates: list[DocumentTemplate],
        context_text: str,
    ) -> dict[str, Any]:
        fallback: dict[str, Any] = {
            "type": None,
            "name_hint": None,
            "subject": None,
            "recipient": None,
            "variables": {},
        }
        try:
            client = self._client()
            completion = await client.chat.completions.create(
                model=LLM_MODEL,
                messages=[
                    {"role": "system", "content": DOCUMENT_EXTRACTION_PROMPT},
                    {
                        "role": "user",
                        "content": build_document_extraction_user_prompt(
                            description,
                            [{"name": t.name, "type": t.type.value} for t in templates],
                            context_text or None,
                        ),
                    },
                ],
                response_format={"type": "json_object"},
            )
            raw = (completion.choices[0].message.content if completion.choices else None) or "{}"
            parsed = json.loads(raw)
            raw_type = parsed.get("type")
            return {
                "type": DocumentTemplateType(raw_type) if raw_type in _TEMPLATE_TYPE_VALUES else None,
                "name_hint": parsed.get("nameHint"),
                "subject": parsed.get("subject"),
                "recipient": parsed.get("recipient"),
                "variables": parsed["variables"]
                if isinstance(parsed.get("variables"), dict)
                else {},
            }
        except Exception as err:
            logger.warning(f"[docgen] extraction failed: {err}")
            return fallback

    async def _generate_grounded(
        self,
        ctx: AuthContext,
        request: dict[str, Any],
        context: dict[str, Any],
    ) -> RagQueryResponse:
        """Free-form document generated straight from retrieved M365 content."""
        client = self._client()
        completion = await client.chat.completions.create(
            model=LLM_MODEL,
            messages=[
                {"role": "system", "content": GROUNDED_DOCUMENT_PROMPT},
                {
                    "role": "user",
                    "content": build_grounded_document_user_prompt(
                        request["description"], context["context_text"]
                    ),
                },
            ],
        )
        generated = (
            (completion.choices[0].message.content if completion.choices else None) or ""
        ).strip()

        await self.documents._log_audit(
            ctx,
            "document.generate.chat",
            "template",
            None,
            {
                "type": "grounded",
                "explicit": request.get("explicit"),
                "grounded": True,
            },
        )

        header = (
            "**Generated document** — drafted from your Microsoft 365 sources. "
            "Review before use.\n\n---\n\n"
        )
        return RagQueryResponse(
            answer=f"{header}{generated}",
            citations=context["citations"],
            chunks=[],
            intent="document_generation",
            document_draft=DocumentDraft(
                type=DocumentTemplateType.report,
                template_name="Generated document",
                body=generated,
                can_save_to_outlook=False,
            ),
        )

    # --- Response builders ----------------------------------------------------

    def _draft_response(
        self,
        template: DocumentTemplate,
        generated: str,
        extracted: dict[str, Any],
        citations: list[Citation],
    ) -> RagQueryResponse:
        is_email = template.type == DocumentTemplateType.customer_email
        subject = (
            extracted["subject"]
            or extracted["variables"].get("subject")
            or template.name
        ) if is_email else None

        source_note = (
            " Values were drawn from your Microsoft 365 sources where available."
            if len(citations) > 0
            else ""
        )
        header = (
            f"**Draft {_label_for_type(template.type)}** — generated from "
            f"*{template.name}*.{source_note}\n\nReview and edit before use; bracketed "
            "values still need your confirmation.\n\n---\n\n"
        )

        return RagQueryResponse(
            answer=f"{header}{generated}",
            citations=citations,
            chunks=[],
            intent="document_generation",
            document_draft=DocumentDraft(
                type=template.type,
                template_name=template.name,
                subject=subject,
                body=generated,
                can_save_to_outlook=is_email,
            ),
        )

    def _help_response(
        self, templates: list[DocumentTemplate], lead: str
    ) -> RagQueryResponse:
        listing = "\n".join(
            f"- **{t.name}** ({_label_for_type(t.type)})" for t in templates
        )
        answer = (
            f"{lead}\n\nYou can ask me to draft a document using one of these "
            f"templates:\n\n{listing}\n\nFor example: *\"Draft an estimate for the "
            "Henderson kitchen remodel\"* or *\"/document write a customer email to "
            "Acme about the delayed delivery\"*. You can also point me at your data, "
            'e.g. *"generate a summary from the SharePoint billing invoice"*.'
        )
        return self._plain_response(answer)

    def _plain_response(self, answer: str) -> RagQueryResponse:
        return RagQueryResponse(
            answer=answer,
            citations=[],
            chunks=[],
            intent="document_generation",
        )


# --- module-level helpers (port of the standalone TS functions) ----------------

_TEMPLATE_TYPE_VALUES = {t.value for t in DocumentTemplateType}


def _detect_sources(description: str) -> list[VectorSource] | None:
    t = description.lower()
    sources: list[VectorSource] = []
    if re.search(r"\bsharepoint\b", t):
        sources.append(VectorSource.sharepoint)
    if re.search(r"\bonedrive\b", t):
        sources.append(VectorSource.onedrive)
    if re.search(r"\b(email|e-mail|mail|inbox|outlook)\b", t):
        sources.append(VectorSource.outlook)
    return sources or None


def _chunk_label(chunk: RetrievedChunk) -> str:
    m = chunk.metadata
    if m.source == VectorSource.outlook:
        return m.subject or "Email"
    return m.file_name or (m.source.value if hasattr(m.source, "value") else str(m.source))


def _label_for_type(type: DocumentTemplateType) -> str:
    return {
        DocumentTemplateType.estimate: "estimate",
        DocumentTemplateType.job_summary: "job summary",
        DocumentTemplateType.customer_email: "customer email",
        DocumentTemplateType.quotation: "quotation",
        DocumentTemplateType.report: "report",
    }.get(type, "document")


# ===========================================================================
# Document export (markdown -> .pdf / .docx)
# ===========================================================================


class _Block:
    """A parsed document block (heading / li / p / hr / table).

    ``table`` blocks (produced by the AI layout path, never by the markdown
    parser) carry ``headers`` and ``rows``; the others use ``text``.
    """

    __slots__ = ("headers", "level", "ordered", "rows", "text", "type")

    def __init__(
        self,
        type: str,
        level: int | None = None,
        ordered: bool | None = None,
        text: str | None = None,
        headers: list[str] | None = None,
        rows: list[list[str]] | None = None,
    ) -> None:
        self.type = type
        self.level = level
        self.ordered = ordered
        self.text = text
        self.headers = headers
        self.rows = rows


_HEADING_RE = re.compile(r"^(#{1,6})\s+(.*)$")
_HR_RE = re.compile(r"^(-{3,}|\*{3,}|_{3,})$")
_UL_RE = re.compile(r"^\s*[-*+]\s+(.*)$")
_OL_RE = re.compile(r"^\s*\d+[.)]\s+(.*)$")


def _parse_blocks(markdown: str) -> list[_Block]:
    lines = markdown.replace("\r\n", "\n").split("\n")
    blocks: list[_Block] = []
    para: list[str] = []

    def flush() -> None:
        if para:
            blocks.append(_Block("p", text=" ".join(para)))
            para.clear()

    for raw in lines:
        line = raw.rstrip()
        if not line.strip():
            flush()
            continue
        m = _HEADING_RE.match(line)
        if m:
            flush()
            blocks.append(_Block("heading", level=len(m.group(1)), text=m.group(2)))
            continue
        if _HR_RE.match(line.strip()):
            flush()
            blocks.append(_Block("hr"))
            continue
        m = _UL_RE.match(line)
        if m:
            flush()
            blocks.append(_Block("li", ordered=False, text=m.group(1)))
            continue
        m = _OL_RE.match(line)
        if m:
            flush()
            blocks.append(_Block("li", ordered=True, text=m.group(1)))
            continue
        para.append(line.strip())

    flush()
    return blocks


_HEADING_SIZE = {1: 20, 2: 16, 3: 14, 4: 12, 5: 11, 6: 11}


class DocumentExportService:
    """Render generated document markdown into downloadable .pdf / .docx bytes.

    Two paths:

      - :meth:`export_ai` (preferred) makes a dedicated LLM call that re-expresses
        the finished markdown as a structured layout (headings, lists, and proper
        **tables**), then renders that. This is what the download buttons use, so
        cost breakdowns / line items come out as real bordered tables rather than
        garbled pipe text.
      - :meth:`export` renders straight from the built-in markdown parser. It is
        the deterministic fallback used when the layout call is unavailable
        (no API key) or fails.

    PDF uses ``reportlab``; DOCX uses ``python-docx``. Both render from a shared
    ``list[_Block]`` so the AI and fallback paths produce identical styling.
    """

    def __init__(self) -> None:
        self._openai: Any | None = None

    # --- OpenAI client (lazy) -------------------------------------------------

    def _client(self) -> Any:
        if self._openai is None:
            if not settings.openai_api_key:
                msg = "OPENAI_API_KEY is not configured"
                raise RuntimeError(msg)
            from openai import AsyncOpenAI

            self._openai = AsyncOpenAI(api_key=settings.openai_api_key)
        return self._openai

    # --- Entry points ---------------------------------------------------------

    async def export(self, markdown: str, format: str) -> bytes:
        """Render straight from the markdown parser (deterministic fallback)."""
        blocks = _parse_blocks(markdown)
        return self._render_docx(blocks) if format == "docx" else self._render_pdf(blocks)

    async def export_ai(self, markdown: str, format: str) -> bytes:
        """Render via a dedicated layout LLM call, falling back to the parser.

        The layout call turns tabular content into real ``table`` blocks; on any
        failure (no key, malformed JSON, empty layout) we fall back to the plain
        markdown parser so the download always succeeds.
        """
        try:
            layout = await self._structure(markdown)
            blocks = _blocks_from_layout(layout)
            if not blocks:
                raise ValueError("layout produced no blocks")
        except Exception as err:
            logger.warning(f"[docexport] AI layout failed, using markdown parser: {err}")
            blocks = _parse_blocks(markdown)
        return self._render_docx(blocks) if format == "docx" else self._render_pdf(blocks)

    async def _structure(self, markdown: str) -> dict[str, Any]:
        """Call the LLM to convert finished markdown into a structured layout."""
        client = self._client()
        completion = await client.chat.completions.create(
            model=LLM_MODEL,
            messages=[
                {"role": "system", "content": DOCUMENT_LAYOUT_PROMPT},
                {"role": "user", "content": build_document_layout_user_prompt(markdown)},
            ],
            response_format={"type": "json_object"},
        )
        raw = (completion.choices[0].message.content if completion.choices else None) or "{}"
        parsed = json.loads(raw)
        return parsed if isinstance(parsed, dict) else {}

    # --- Renderers (shared by both paths) -------------------------------------

    def _render_docx(self, blocks: list[_Block]) -> bytes:
        from io import BytesIO

        from docx import Document
        from docx.shared import Pt

        doc = Document()
        normal = doc.styles["Normal"]
        normal.font.name = "Calibri"
        normal.font.size = Pt(11)

        for b in blocks:
            if b.type == "heading":
                doc.add_heading(_strip_bold(b.text or ""), level=min(b.level or 3, 9))
            elif b.type == "hr":
                doc.add_paragraph("_" * 40)
            elif b.type == "table":
                _add_docx_table(doc, b.headers or [], b.rows or [])
            elif b.type == "li":
                style = "List Number" if b.ordered else "List Bullet"
                p = doc.add_paragraph(style=style)
                _add_inline_runs(p, b.text or "")
            else:
                p = doc.add_paragraph()
                _add_inline_runs(p, b.text or "")

        buffer = BytesIO()
        doc.save(buffer)
        return buffer.getvalue()

    def _render_pdf(self, blocks: list[_Block]) -> bytes:
        from io import BytesIO

        from reportlab.lib.enums import TA_LEFT
        from reportlab.lib.pagesizes import A4
        from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
        from reportlab.lib.units import mm
        from reportlab.platypus import (
            HRFlowable,
            ListFlowable,
            ListItem,
            Paragraph,
            SimpleDocTemplate,
            Spacer,
        )

        buffer = BytesIO()
        doc = SimpleDocTemplate(
            buffer,
            pagesize=A4,
            leftMargin=20 * mm,
            rightMargin=20 * mm,
            topMargin=20 * mm,
            bottomMargin=20 * mm,
        )
        styles = getSampleStyleSheet()
        body_style = ParagraphStyle(
            "DocBody",
            parent=styles["Normal"],
            fontName="Helvetica",
            fontSize=11,
            leading=15,
            alignment=TA_LEFT,
        )

        flowables: list[Any] = []
        for b in blocks:
            if b.type == "heading":
                size = _HEADING_SIZE.get(b.level or 3, 14)
                heading_style = ParagraphStyle(
                    f"DocHeading{b.level}",
                    parent=body_style,
                    fontName="Helvetica-Bold",
                    fontSize=size,
                    leading=size + 4,
                    spaceBefore=6,
                    spaceAfter=4,
                )
                flowables.append(Paragraph(_inline_to_pdf(b.text or ""), heading_style))
            elif b.type == "hr":
                flowables.append(Spacer(1, 4))
                flowables.append(HRFlowable(width="100%", color="#cccccc"))
                flowables.append(Spacer(1, 6))
            elif b.type == "table":
                flowables.append(_build_pdf_table(b.headers or [], b.rows or [], body_style))
                flowables.append(Spacer(1, 8))
            elif b.type == "li":
                item = ListItem(
                    Paragraph(_inline_to_pdf(b.text or ""), body_style),
                    leftIndent=12,
                )
                flowables.append(
                    ListFlowable(
                        [item],
                        bulletType="1" if b.ordered else "bullet",
                    )
                )
            else:
                flowables.append(Paragraph(_inline_to_pdf(b.text or ""), body_style))
                flowables.append(Spacer(1, 6))

        doc.build(flowables)
        return buffer.getvalue()


def _strip_bold(text: str) -> str:
    return text.replace("**", "")


def _escape_html(s: str) -> str:
    return s.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _inline_to_pdf(text: str) -> str:
    """Convert **bold** markers to reportlab markup, escaping the rest."""
    escaped = _escape_html(text)
    return re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", escaped)


def _add_inline_runs(paragraph: Any, text: str) -> None:
    """Add **bold**-aware runs to a python-docx paragraph."""
    parts = text.split("**")
    for i, part in enumerate(parts):
        if part == "":
            continue
        run = paragraph.add_run(part)
        if i % 2 == 1:
            run.bold = True


# --- AI layout -> blocks --------------------------------------------------------

# JSON block "type" (from DOCUMENT_LAYOUT_PROMPT) -> internal _Block list builder.
def _blocks_from_layout(layout: dict[str, Any]) -> list[_Block]:
    """Convert the LLM's structured-layout JSON into the internal block list.

    Tolerant of partial / slightly-off shapes (the model occasionally omits a
    field): unknown block types and empty entries are skipped rather than raised.
    """
    blocks: list[_Block] = []
    for raw in layout.get("blocks") or []:
        if not isinstance(raw, dict):
            continue
        kind = raw.get("type")
        if kind == "heading":
            text = str(raw.get("text") or "").strip()
            if text:
                blocks.append(_Block("heading", level=_clamp_level(raw.get("level")), text=text))
        elif kind == "paragraph":
            text = str(raw.get("text") or "").strip()
            if text:
                blocks.append(_Block("p", text=text))
        elif kind in ("bullets", "numbers"):
            ordered = kind == "numbers"
            for item in raw.get("items") or []:
                item_text = str(item).strip()
                if item_text:
                    blocks.append(_Block("li", ordered=ordered, text=item_text))
        elif kind == "table":
            headers = [str(h) for h in (raw.get("headers") or [])]
            rows = [
                [str(c) for c in row]
                for row in (raw.get("rows") or [])
                if isinstance(row, list)
            ]
            if headers or rows:
                blocks.append(_Block("table", headers=headers, rows=rows))
        elif kind == "divider":
            blocks.append(_Block("hr"))

    # Prepend the document title as a top-level heading, unless the first block
    # is already that same heading (the prompt asks the model not to repeat it,
    # but we guard against it anyway).
    title = str(layout.get("title") or "").strip()
    if title:
        first = blocks[0] if blocks else None
        already = (
            first is not None
            and first.type == "heading"
            and (first.text or "").strip().lower() == title.lower()
        )
        if not already:
            blocks.insert(0, _Block("heading", level=1, text=title))
    return blocks


def _clamp_level(level: Any) -> int:
    try:
        return max(1, min(6, int(level)))
    except (TypeError, ValueError):
        return 2


def _add_docx_table(doc: Any, headers: list[str], rows: list[list[str]]) -> None:
    """Append a bordered table (bold header row) to a python-docx document."""
    ncols = len(headers) or (len(rows[0]) if rows else 0)
    if ncols == 0:
        return

    table = doc.add_table(rows=0, cols=ncols)
    try:
        table.style = "Table Grid"  # built-in; gives every cell a border
    except KeyError:  # pragma: no cover - style always ships with python-docx
        pass

    if headers:
        cells = table.add_row().cells
        for i in range(ncols):
            para = cells[i].paragraphs[0]
            _add_inline_runs(para, headers[i] if i < len(headers) else "")
            for run in para.runs:
                run.bold = True

    for row in rows:
        cells = table.add_row().cells
        for i in range(ncols):
            _add_inline_runs(cells[i].paragraphs[0], row[i] if i < len(row) else "")

    doc.add_paragraph()  # breathing room after the table


def _build_pdf_table(headers: list[str], rows: list[list[str]], body_style: Any) -> Any:
    """Build a reportlab ``Table`` flowable (grid + shaded, bold header row)."""
    from reportlab.lib import colors
    from reportlab.lib.styles import ParagraphStyle
    from reportlab.platypus import Paragraph, Table, TableStyle

    ncols = len(headers) or (len(rows[0]) if rows else 1)
    header_style = ParagraphStyle(
        "DocTableHeader", parent=body_style, fontName="Helvetica-Bold", fontSize=10, leading=13
    )
    cell_style = ParagraphStyle(
        "DocTableCell", parent=body_style, fontSize=10, leading=13
    )

    def _pad(values: list[str]) -> list[str]:
        return [values[i] if i < len(values) else "" for i in range(ncols)]

    data: list[list[Any]] = []
    if headers:
        data.append([Paragraph(_inline_to_pdf(c), header_style) for c in _pad(headers)])
    for row in rows:
        data.append([Paragraph(_inline_to_pdf(c), cell_style) for c in _pad(row)])
    if not data:  # nothing to render — emit an empty single cell to stay valid
        data = [[Paragraph("", cell_style)]]

    table = Table(data, hAlign="LEFT")
    style: list[Any] = [
        ("GRID", (0, 0), (-1, -1), 0.5, colors.HexColor("#cccccc")),
        ("VALIGN", (0, 0), (-1, -1), "TOP"),
        ("LEFTPADDING", (0, 0), (-1, -1), 6),
        ("RIGHTPADDING", (0, 0), (-1, -1), 6),
        ("TOPPADDING", (0, 0), (-1, -1), 4),
        ("BOTTOMPADDING", (0, 0), (-1, -1), 4),
    ]
    if headers:
        style.append(("BACKGROUND", (0, 0), (-1, 0), colors.HexColor("#f0f0f0")))
    table.setStyle(TableStyle(style))
    return table


# Typed-dict-style aliases used in method signatures above (kept as plain dicts at
# runtime; documented here for readers). These mirror the NestJS input interfaces.
CreateTemplateInput = dict[str, Any]
UpdateTemplateInput = dict[str, Any]
GenerateDocumentRequestInput = dict[str, Any]
